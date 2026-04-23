from __future__ import annotations

"""이벤트 기록 CSV와 문서 템플릿 생성을 담당하는 공통 로직."""

import json
import os
from dataclasses import dataclass
from pathlib import Path

import pandas as pd
from docx import Document


SHORT_THRESHOLD_SECONDS = 5.0


@dataclass(frozen=True)
class GenerationContext:
    """LLM 사용 가능 여부와 사유를 함께 담는 설정 객체."""
    llm_enabled: bool
    reason: str


def load_input_events(input_path: Path) -> pd.DataFrame:
    """입력 엑셀을 읽고 시간/이벤트 컬럼을 정규화한다."""
    frame = pd.read_excel(input_path, header=0)
    required_columns = {"time_sec", "flag_id"}
    missing_columns = required_columns - set(frame.columns)
    if missing_columns:
        raise ValueError(f"Missing required columns: {sorted(missing_columns)}")

    frame = frame.copy()
    frame["time_sec"] = pd.to_datetime(frame["time_sec"])
    frame["flag_norm"] = frame["flag_id"].apply(normalize_flag)
    frame = frame.sort_values("time_sec").reset_index(drop=True)
    return frame


def normalize_flag(raw_value: object) -> str | None:
    """원본 이벤트 값을 A/S/D 표준 플래그로 정규화한다."""
    if pd.isna(raw_value):
        return None
    text = str(raw_value).strip().upper()
    if text in {"1", "A"}:
        return "A"
    if text in {"2", "S"}:
        return "S"
    if text in {"3", "D"}:
        return "D"
    return None


def format_time(value: pd.Timestamp) -> str:
    """문서 출력용 시각 문자열을 만든다."""
    return value.strftime("%H:%M:%S")


def build_event_segments(frame: pd.DataFrame) -> pd.DataFrame:
    """프레임 단위 이벤트를 연속 구간 단위 이벤트로 합친다."""
    events: list[dict[str, object]] = []
    current_flag: str | None = None
    current_start: pd.Timestamp | None = None
    previous_time: pd.Timestamp | None = None
    event_index = 0

    for _, row in frame.iterrows():
        current_time = row["time_sec"]
        current_value = row["flag_norm"]

        if current_value not in {"A", "S", "D"}:
            if current_flag is not None and current_start is not None and previous_time is not None:
                # 유효한 A/S/D 구간이 끝난 시점에 이벤트를 확정한다.
                events.append(
                    {
                        "idx": event_index,
                        "flag": current_flag,
                        "start_time": current_start,
                        "end_time": previous_time,
                        "duration_sec": float((previous_time - current_start).total_seconds()),
                    }
                )
                event_index += 1
                current_flag = None
                current_start = None
            previous_time = current_time
            continue

        if current_flag is None:
            current_flag = current_value
            current_start = current_time
        elif current_value != current_flag and current_start is not None and previous_time is not None:
            events.append(
                {
                    "idx": event_index,
                    "flag": current_flag,
                    "start_time": current_start,
                    "end_time": previous_time,
                    "duration_sec": float((previous_time - current_start).total_seconds()),
                }
            )
            event_index += 1
            current_flag = current_value
            current_start = current_time

        previous_time = current_time

    if current_flag is not None and current_start is not None and previous_time is not None:
        events.append(
            {
                "idx": event_index,
                "flag": current_flag,
                "start_time": current_start,
                "end_time": previous_time,
                "duration_sec": float((previous_time - current_start).total_seconds()),
            }
        )

    return pd.DataFrame(events)


def build_auto_log(events_df: pd.DataFrame) -> tuple[pd.DataFrame, set[int], set[int], list[dict[str, object]]]:
    """이벤트 구간을 자동 기록표 형식으로 변환하고 이상 후보도 함께 추린다."""
    short_event_indices = set(events_df.loc[events_df["duration_sec"] <= SHORT_THRESHOLD_SECONDS, "idx"].tolist())
    missing_process_pairs: list[dict[str, object]] = []

    for index, event in events_df.iterrows():
        if event["flag"] != "A":
            continue
        next_index = index + 1
        while next_index < len(events_df) and events_df.iloc[next_index]["flag"] == "A":
            next_index += 1
        if next_index >= len(events_df):
            continue
        next_event = events_df.iloc[next_index]
        if next_event["flag"] == "D":
            missing_process_pairs.append(
                {
                    "a_event": event.to_dict(),
                    "d_event": next_event.to_dict(),
                }
            )

    missing_process_indices = {int(pair["d_event"]["idx"]) for pair in missing_process_pairs}

    rows: list[dict[str, str]] = []
    for _, event in events_df.iterrows():
        event_index = int(event["idx"])
        remarks = ""
        status = "정상"
        if event_index in short_event_indices:
            status = "확인 필요"
        if event_index in missing_process_indices:
            remarks = "작업 프로세스 확인 필요"
        rows.append(
            {
                "시간(Time)": f"{format_time(event['start_time'])} ~ {format_time(event['end_time'])}",
                "감지된 행동(AI Event)": event["flag"],
                "적합 여부(Status)": status,
                "비고(Remarks)": remarks,
            }
        )

    return pd.DataFrame(rows), short_event_indices, missing_process_indices, missing_process_pairs


def resolve_generation_context() -> GenerationContext:
    """Gemini API 사용 가능 여부를 점검해 문장 생성 모드를 결정한다."""
    api_key = os.getenv("GEMINI_API_KEY", "").strip()
    if not api_key:
        return GenerationContext(llm_enabled=False, reason="GEMINI_API_KEY is not set")
    try:
        import google.generativeai as genai

        genai.configure(api_key=api_key)
        return GenerationContext(llm_enabled=True, reason="Gemini enabled")
    except Exception as error:
        return GenerationContext(llm_enabled=False, reason=f"Gemini import/config failed: {error}")


def generate_anomaly_text(item: dict[str, object], context: GenerationContext) -> tuple[str, str]:
    """이상 이벤트 설명과 조치 문구를 생성한다."""
    if context.llm_enabled:
        try:
            import google.generativeai as genai

            model = genai.GenerativeModel("gemini-1.5-pro")
            prompt = build_anomaly_prompt(json.dumps(item, ensure_ascii=False))
            response = model.generate_content(prompt)
            parsed = json.loads(response.text.strip())
            description = parsed.get("description", "").strip()
            action = parsed.get("action", "").strip()
            if description and action:
                return description, action
        except Exception:
            pass

    # LLM을 쓰지 못할 때도 문서 생성이 멈추지 않도록 기본 문구를 제공한다.
    if item["type"] == "short_duration":
        description = (
            f"{item['flag']} 동작이 {item['start_time_str']}부터 {item['end_time_str']}까지 "
            f"{item['duration_sec']:.3f}초 동안만 감지되었습니다."
        )
        action = "해당 시간대 작업 로그와 영상 기록을 확인합니다."
    else:
        description = "A 이후 중간 공정 없이 D로 전환되어 작업 순서 누락 가능성이 확인되었습니다."
        action = "해당 배치의 작업 프로세스와 현장 기록을 확인합니다."
    return description, action


def build_anomaly_prompt(event_json: str) -> str:
    """LLM에 전달할 이상 이벤트 요약 프롬프트를 만든다."""
    return f"""
당신은 제조 공정 기록 담당자입니다.
다음 이상 이벤트 정보를 바탕으로 기록지에 들어갈 '상세 내용'과 '자동 조치'를 한국어 JSON으로 작성하세요.

입력:
{event_json}

출력 형식:
{{
  "description": "<상세 내용>",
  "action": "<자동 조치>"
}}
"""


def build_anomaly_log(
    events_df: pd.DataFrame,
    short_event_indices: set[int],
    missing_process_pairs: list[dict[str, object]],
    context: GenerationContext,
) -> pd.DataFrame:
    """이상 후보를 문서용 이상 기록표 형식으로 변환한다."""
    anomaly_items: list[dict[str, object]] = []

    for _, event in events_df.iterrows():
        event_index = int(event["idx"])
        if event_index not in short_event_indices:
            continue
        anomaly_items.append(
            {
                "type": "short_duration",
                "flag": event["flag"],
                "start_time_str": format_time(event["start_time"]),
                "end_time_str": format_time(event["end_time"]),
                "duration_sec": float(event["duration_sec"]),
            }
        )

    for pair in missing_process_pairs:
        anomaly_items.append(
            {
                "type": "missing_process",
                "prev_flag": pair["a_event"]["flag"],
                "next_flag": pair["d_event"]["flag"],
                "prev_start_time_str": format_time(pair["a_event"]["start_time"]),
                "prev_end_time_str": format_time(pair["a_event"]["end_time"]),
                "next_start_time_str": format_time(pair["d_event"]["start_time"]),
                "next_end_time_str": format_time(pair["d_event"]["end_time"]),
            }
        )

    rows: list[dict[str, str]] = []
    for item in anomaly_items:
        if item["type"] == "short_duration":
            time_text = f"{item['start_time_str']} ~ {item['end_time_str']}"
            event_type = "5초 이하 동작"
        else:
            time_text = f"{item['prev_start_time_str']} ~ {item['next_end_time_str']}"
            event_type = "행동 순서 누락(A-중간-D)"
        description, action = generate_anomaly_text(item, context)
        rows.append(
            {
                "시간(Time)": time_text,
                "이상유형(Event Type)": event_type,
                "상세 내용(Description)": description,
                "자동 조치(Action)": action,
                "담당자 확인(Check)": "",
            }
        )

    return pd.DataFrame(rows)


def find_table_index_by_keyword(document: Document, keyword: str) -> int | None:
    """키워드가 포함된 표를 찾아 템플릿 내 인덱스를 반환한다."""
    for table_index, table in enumerate(document.tables):
        for row in table.rows:
            row_text = " ".join(cell.text for cell in row.cells)
            if keyword in row_text:
                return table_index
    return None


def _clear_table_rows(table, keep_header_rows: int = 1) -> None:
    """헤더를 제외한 기존 행을 모두 지운다."""
    while len(table.rows) > keep_header_rows:
        table._tbl.remove(table.rows[-1]._tr)


def append_rows_to_table(table, rows: list[dict[str, str]], columns: list[tuple[int, str]]) -> None:
    """행 딕셔너리 목록을 워드 표에 순서대로 채운다."""
    _clear_table_rows(table, keep_header_rows=1)
    for row_data in rows:
        row = table.add_row().cells
        for column_index, column_name in columns:
            row[column_index].text = str(row_data.get(column_name, ""))


def fill_template_document(
    template_path: Path,
    auto_log_df: pd.DataFrame,
    anomaly_df: pd.DataFrame,
    output_path: Path,
) -> Path:
    """자동 기록표와 이상 기록표를 DOCX 템플릿에 채워 저장한다."""
    document = Document(template_path)

    auto_table_index = find_table_index_by_keyword(document, "감지된 행동")
    if auto_table_index is None:
        auto_table_index = 2 if len(document.tables) > 2 else 0

    anomaly_table_index = find_table_index_by_keyword(document, "이상유형")
    if anomaly_table_index is None:
        anomaly_table_index = 3 if len(document.tables) > 3 else min(1, len(document.tables) - 1)

    append_rows_to_table(
        document.tables[auto_table_index],
        auto_log_df.to_dict("records"),
        [
            (0, "시간(Time)"),
            (1, "감지된 행동(AI Event)"),
            (2, "적합 여부(Status)"),
            (3, "비고(Remarks)"),
        ],
    )
    append_rows_to_table(
        document.tables[anomaly_table_index],
        anomaly_df.to_dict("records"),
        [
            (0, "시간(Time)"),
            (1, "이상유형(Event Type)"),
            (2, "상세 내용(Description)"),
            (3, "자동 조치(Action)"),
            (4, "담당자 확인(Check)"),
        ],
    )

    document.save(output_path)
    return output_path
